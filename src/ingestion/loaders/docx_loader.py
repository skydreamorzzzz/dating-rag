"""Loader for DOCX files via python-docx."""

from pathlib import Path
from docx import Document as DocxDocument

from src.ingestion.loaders.base import BaseLoader
from src.ingestion.schemas import Document
from src.utils.logger import get_logger

logger = get_logger(__name__)


class DocxLoader(BaseLoader):
    """Load a .docx file, extracting paragraphs and table text into one Document."""

    def load(self, file_path: Path) -> list[Document]:
        try:
            docx = DocxDocument(str(file_path))
        except Exception as e:
            raise ValueError(f"Failed to open DOCX {file_path.name}: {e}") from e

        parts: list[str] = []

        # --- paragraphs ---
        for para in docx.paragraphs:
            text = para.text
            if text:
                parts.append(text)

        # --- tables ---
        for table in docx.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                rows.append("\t".join(cells))
            if rows:
                parts.append("\n".join(rows))

        full_text = "\n".join(parts)
        total_chars = len(full_text)

        warnings: list[str] = []
        if not full_text.strip():
            warnings.append("No text extracted from DOCX — file may be empty or image-only")

        doc = Document(
            source_file=file_path.name,
            source_path=str(file_path),
            doc_type="docx",
            text=full_text,
            warnings=warnings,
        )
        logger.info(
            "Loaded %s (%d chars, %d paragraphs, %d tables)",
            file_path.name, total_chars, len(docx.paragraphs), len(docx.tables),
        )
        return [doc]

    @staticmethod
    def supported_extensions() -> list[str]:
        return [".docx"]
